from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "请填写学号+姓名"
DOCX_PATH = OUT_DIR / "软件开发文档.docx"
SCREENSHOT_DIR = ROOT / "docs" / "screenshots"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header=True):
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
            if row_index == 0 and header:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_heading(doc, text, level):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        if level == 1:
            run.font.color.rgb = RGBColor(46, 116, 181)
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = RGBColor(46, 116, 181)
            run.font.size = Pt(13)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10.5)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)
    return paragraph


def add_screenshot(doc, filename, caption):
    path = SCREENSHOT_DIR / filename
    if path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(6.2))
        add_caption(doc, caption)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(4.75)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "项目", True)
    set_cell_text(hdr[1], "内容", True)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
    style_table(table)
    doc.add_paragraph()


def add_three_col_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1


def build():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("CourseHub 课程选课管理系统")
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("JavaWeb 期末项目软件开发文档")
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(85, 85, 85)

    add_key_value_table(doc, [
        ("课程名称", "JavaWeb 开发技术"),
        ("项目名称", "CourseHub 课程选课管理系统"),
        ("项目类型", "B/S 架构动态 Web 应用"),
        ("开发技术", "JSP、Servlet、JavaBean、JDBC、Filter、Listener、Bootstrap、Ajax"),
        ("学生信息", "请填写：姓名、学号"),
        ("提交日期", "2026 年 7 月 1 日前"),
    ])

    add_heading(doc, "1. 项目概述", 1)
    add_body(doc, "CourseHub 是一个面向课程选课场景的 JavaWeb 课程项目，采用 B/S 架构实现管理员课程管理、学生在线选课、附件上传下载、数据导出和后台用户管理等功能。系统遵循 MVC 分层思想，使用 Servlet 作为控制层，Service 处理业务逻辑，DAO 通过 JDBC 操作 MySQL 数据库，JSP 负责页面展示。")
    add_body(doc, "本项目围绕课程期末任务书要求设计，覆盖 JSP、Servlet、JavaBean、Session、Cookie、JDBC、Filter、Ajax、文件上传下载、分页、搜索、事务和权限控制等核心知识点。")

    add_heading(doc, "2. 需求分析", 1)
    add_bullet(doc, "用户模块：支持注册、登录、注销、头像上传和个人资料维护。")
    add_bullet(doc, "权限模块：系统包含管理员和普通用户两种角色，未登录用户访问受保护页面时自动跳转登录页。")
    add_bullet(doc, "课程模块：管理员可新增、修改、删除、查询课程，支持课程附件与课程图片上传。")
    add_bullet(doc, "选课模块：普通用户可选择课程、上传多个附件、修改备注和退课；管理员可查看全部选课记录。")
    add_bullet(doc, "查询与导出：课程和选课列表支持条件搜索、分页、排序，选课记录支持 CSV 导出。")
    add_bullet(doc, "安全要求：密码使用 SHA-256 摘要保存，所有 SQL 使用 PreparedStatement，表单携带 CSRF Token，显示内容进行 HTML 转义。")

    add_heading(doc, "3. 系统总体设计", 1)
    add_body(doc, "系统采用典型 JavaWeb MVC 分层结构：JSP 页面负责视图展示，Servlet 接收请求并调用 Service，Service 完成业务校验和事务处理，DAO 封装 JDBC 数据访问逻辑，实体类作为 JavaBean 在各层之间传递数据。")
    add_three_col_table(doc, ["层次", "主要包/文件", "职责"], [
        ("Controller", "com.hnit.coursehub.controller", "接收请求、参数处理、页面跳转、调用 Service"),
        ("Service", "com.hnit.coursehub.service", "业务校验、权限相关判断、事务管理"),
        ("DAO", "com.hnit.coursehub.dao", "通过 PreparedStatement 执行 SQL，封装数据库访问"),
        ("Entity", "com.hnit.coursehub.entity", "JavaBean 实体类，实现 Serializable"),
        ("Filter/Listener", "filter、listener", "编码过滤、登录拦截、管理员拦截、CSRF 校验、在线人数统计"),
        ("View", "src/main/webapp/WEB-INF/views", "JSP + JSTL + EL 页面展示"),
    ])

    add_heading(doc, "4. 数据库设计", 1)
    add_body(doc, "数据库名称为 coursehub，核心表包括 users、courses、enrollments、enrollment_files。表之间通过外键关联，选课记录表同时关联用户和课程，实现多表 JOIN 查询。")
    add_three_col_table(doc, ["表名", "说明", "关键字段"], [
        ("users", "用户表", "id、username、password_hash、real_name、email、role、avatar_path"),
        ("courses", "课程表", "id、course_code、course_name、teacher、category、capacity、created_by"),
        ("enrollments", "选课记录表", "id、user_id、course_id、status、note、created_at"),
        ("enrollment_files", "选课附件表", "id、enrollment_id、original_name、file_path、file_size"),
    ])
    add_body(doc, "选课业务使用 JDBC 事务：新增选课时同时写入 enrollments、enrollment_files 并扣减课程名额；退课时删除选课记录并恢复课程名额；删除用户或课程时清理关联数据，保证数据一致性。")

    add_heading(doc, "5. 功能实现说明", 1)
    add_heading(doc, "5.1 用户与权限", 2)
    add_body(doc, "用户登录成功后，系统将登录用户对象保存到 Session；同时通过 Cookie 记住用户名。AuthFilter 负责登录拦截，AdminFilter 负责管理员权限拦截。")
    add_heading(doc, "5.2 课程管理", 2)
    add_body(doc, "管理员可维护课程基础信息，包括课程编号、课程名称、教师、分类、剩余名额、课程简介、课程图片和课程附件。列表支持关键词搜索、分类筛选、分页和排序。")
    add_heading(doc, "5.3 选课管理", 2)
    add_body(doc, "普通用户可以在线选课并上传多个附件，管理员可以查看所有用户选课记录。选课列表通过 SQL JOIN 关联用户表和课程表，展示学生、课程、教师、分类、状态和附件下载入口。")
    add_heading(doc, "5.4 文件上传下载与导出", 2)
    add_body(doc, "系统使用 Servlet 3.0 MultipartConfig 处理头像、课程附件和选课附件上传；下载功能通过 OutputStream 输出服务器文件；选课记录导出为 CSV 文件。")

    add_heading(doc, "6. 运行部署说明", 1)
    add_number(doc, "启动 MySQL，并执行 database/init.sql 初始化数据库。")
    add_number(doc, "检查 src/main/resources/db.properties 中的 jdbc.username 和 jdbc.password 是否与本机 MySQL 一致。")
    add_number(doc, "在项目目录执行 mvn clean package，生成 target/coursehub.war。")
    add_number(doc, "将 coursehub.war 复制到 Tomcat 9 的 webapps 目录。")
    add_number(doc, "启动 Tomcat，访问 http://localhost:8080/coursehub。")
    add_three_col_table(doc, ["账号类型", "用户名", "密码"], [
        ("管理员", "admin", "123456"),
        ("普通用户", "student", "123456"),
        ("普通用户", "lisi", "123456"),
    ])

    add_heading(doc, "7. 系统界面截图", 1)
    add_screenshot(doc, "01-login.png", "图 1 登录页面：用户输入账号密码进入系统")
    add_screenshot(doc, "02-dashboard.png", "图 2 管理员仪表盘：展示课程、选课和在线会话概览")
    add_screenshot(doc, "03-course-list.png", "图 3 课程列表：支持搜索、分页、排序和管理员操作")
    add_screenshot(doc, "04-course-form.png", "图 4 课程表单：支持课程信息维护和单文件上传")
    add_screenshot(doc, "05-enrollment-list.png", "图 5 选课记录：多表关联查询、附件下载和 CSV 导出")
    add_screenshot(doc, "06-user-admin.png", "图 6 用户管理：管理员维护用户资料和角色")
    add_screenshot(doc, "07-profile.png", "图 7 个人资料：用户修改资料并上传头像")

    add_heading(doc, "8. 项目总结", 1)
    add_body(doc, "本项目完整实践了 JavaWeb 基础技术栈，从数据库表设计、JDBC 数据访问、Service 业务分层、Servlet 控制器、Filter 权限控制到 JSP 页面展示均有实现。系统功能覆盖课程任务书中的用户注册登录、权限控制、多业务模块 CRUD、多表 JOIN、分页搜索、事务、上传下载、Ajax 交互和文档说明等要求。")
    add_body(doc, "后续可以继续扩展课程审核流程、选课时间段控制、操作日志、Excel 导出和更完整的单元测试。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("CourseHub 课程选课管理系统软件开发文档")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build())
